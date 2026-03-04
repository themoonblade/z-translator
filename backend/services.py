import os
import httpx
import io
from fastapi import UploadFile
from pypdf import PdfReader
import docx
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT
from pdf2docx import Converter
from PIL import Image
import base64

OLLAMA_BASE_URL = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")

class OllamaService:
    async def get_models(self):
        async with httpx.AsyncClient() as client:
            try:
                response = await client.get(f"{OLLAMA_BASE_URL}/api/tags")
                response.raise_for_status()
                return response.json()
            except Exception as e:
                print(f"Error fetching models: {e}")
                return {"models": []}

    async def translate(self, text: str, source_lang: str, target_lang: str, model: str = "llama3", timeout: int = 120):
        import time
        
        start_time = time.time()
        
        # Handle auto-detection
        if source_lang.lower() == "auto":
            prompt = f"""Detect the language of the following text and translate it to {target_lang}.
IMPORTANT: The text contains special markers like ##IMG001##, ##TBL001## etc. These are placeholders for images and tables.
You MUST preserve these markers EXACTLY as they appear in your translation. Do not translate, modify, or remove them.

Text to translate:
{text}

Return ONLY the translated text with the markers preserved."""
        else:
            prompt = f"""Translate the following text from {source_lang} to {target_lang}.
IMPORTANT: The text contains special markers like ##IMG001##, ##TBL001## etc. These are placeholders for images and tables.
You MUST preserve these markers EXACTLY as they appear in your translation. Do not translate, modify, or remove them.

Text to translate:
{text}

Return ONLY the translated text with the markers preserved."""
        
        async with httpx.AsyncClient(timeout=float(timeout)) as client:
            try:
                # Using the /api/generate endpoint for generic models
                response = await client.post(
                    f"{OLLAMA_BASE_URL}/api/generate",
                    json={
                        "model": model,
                        "prompt": prompt,
                        "stream": False
                    }
                )
                response.raise_for_status()
                result = response.json()
                translation = result.get("response", "").strip()
                
                # Calculate metrics
                end_time = time.time()
                duration = end_time - start_time
                word_count = len(text.split())
                words_per_second = word_count / duration if duration > 0 else 0
                
                return {
                    "translation": translation,
                    "duration": round(duration, 2),
                    "word_count": word_count,
                    "words_per_second": round(words_per_second, 2)
                }
            except Exception as e:
                print(f"Error translating text: {e}")
                return {
                    "translation": f"Error: {str(e)}",
                    "duration": 0,
                    "word_count": 0,
                    "words_per_second": 0
                }

    async def summarize(self, text: str, model: str = "llama3", target_lang: str = "English", timeout: int = 120):
        import time
        
        start_time = time.time()
        
        prompt = f"""Summarize the following text concisely in {target_lang}.
IMPORTANT: The text contains special markers like ##IMG001##, ##TBL001## etc. These are placeholders for images and tables.
You MUST preserve these markers EXACTLY in your summary. Do not translate, modify, or remove them.

Text to summarize:
{text}

Return ONLY the summary with the markers preserved."""
        
        async with httpx.AsyncClient(timeout=float(timeout)) as client:
            try:
                response = await client.post(
                    f"{OLLAMA_BASE_URL}/api/generate",
                    json={
                        "model": model,
                        "prompt": prompt,
                        "stream": False
                    }
                )
                response.raise_for_status()
                result = response.json()
                summary = result.get("response", "").strip()
                
                # Calculate metrics
                end_time = time.time()
                duration = end_time - start_time
                word_count = len(text.split())
                summary_word_count = len(summary.split())
                compression_ratio = round((1 - summary_word_count / word_count) * 100, 1) if word_count > 0 else 0
                
                return {
                    "summary": summary,
                    "duration": round(duration, 2),
                    "original_word_count": word_count,
                    "summary_word_count": summary_word_count,
                    "compression_ratio": compression_ratio,
                    "words_per_second": round(word_count / duration, 2) if duration > 0 else 0
                }
            except Exception as e:
                print(f"Error summarizing text: {e}")
                return {
                    "summary": f"Error: {str(e)}",
                    "duration": 0,
                    "original_word_count": 0,
                    "summary_word_count": 0,
                    "compression_ratio": 0,
                    "words_per_second": 0
                }

    async def rewrite(self, text: str, model: str = "llama3", target_lang: str = "English", timeout: int = 120):
        import time
        
        start_time = time.time()
        
        prompt = f"""Rewrite the following text to improve clarity and professionalism in {target_lang}.
IMPORTANT: The text contains special markers like ##IMG001##, ##TBL001## etc. These are placeholders for images and tables.
You MUST preserve these markers EXACTLY in your rewrite. Do not translate, modify, or remove them.

Text to rewrite:
{text}

Return ONLY the rewritten text with the markers preserved."""
        
        async with httpx.AsyncClient(timeout=float(timeout)) as client:
            try:
                response = await client.post(
                    f"{OLLAMA_BASE_URL}/api/generate",
                    json={
                        "model": model,
                        "prompt": prompt,
                        "stream": False
                    }
                )
                response.raise_for_status()
                result = response.json()
                rewritten = result.get("response", "").strip()
                
                # Calculate metrics
                end_time = time.time()
                duration = end_time - start_time
                word_count = len(text.split())
                rewritten_word_count = len(rewritten.split())
                words_per_second = word_count / duration if duration > 0 else 0
                
                return {
                    "rewritten": rewritten,
                    "duration": round(duration, 2),
                    "original_word_count": word_count,
                    "rewritten_word_count": rewritten_word_count,
                    "words_per_second": round(words_per_second, 2)
                }
            except Exception as e:
                print(f"Error rewriting text: {e}")
                return {
                    "rewritten": f"Error: {str(e)}",
                    "duration": 0,
                    "original_word_count": 0,
                    "rewritten_word_count": 0,
                    "words_per_second": 0
                }

class FileParser:
    def __init__(self):
        self.objects = {}  # Store extracted objects with anchors
    
    async def parse_file(self, file: UploadFile) -> str:
        content = await file.read()
        filename = file.filename.lower()
        
        if filename.endswith(".pdf"):
            text, objects = self._parse_pdf(content)
            self.objects = objects
            return text
        elif filename.endswith(".docx"):
            text, objects = self._parse_docx(content)
            self.objects = objects
            return text
        elif filename.endswith(".txt"):
            self.objects = {}
            return content.decode("utf-8")
        elif filename.endswith(".html"):
            self.objects = {}
            return self._parse_html(content)
        else:
            raise ValueError("Unsupported file type")

    def _parse_pdf(self, content: bytes) -> tuple[str, dict]:
        """Convert PDF to DOCX and parse it"""
        # Save PDF temporarily
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as pdf_file:
            pdf_file.write(content)
            pdf_path = pdf_file.name
        
        try:
            # Convert PDF to DOCX
            docx_path = pdf_path.replace('.pdf', '.docx')
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            
            # Parse the converted DOCX
            with open(docx_path, 'rb') as f:
                docx_content = f.read()
            
            text, objects = self._parse_docx(docx_content)
            
            # Cleanup
            import os
            os.unlink(pdf_path)
            os.unlink(docx_path)
            
            return text, objects
        except Exception as e:
            print(f"Error converting PDF: {e}")
            # Fallback to simple text extraction
            import os
            os.unlink(pdf_path)
            reader = PdfReader(io.BytesIO(content))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text, {}

    def _parse_docx(self, content: bytes) -> tuple[str, dict]:
        """Parse DOCX and extract text with anchors for images and tables"""
        doc = docx.Document(io.BytesIO(content))
        text_parts = []
        objects = {}  # {anchor: object_data}
        img_counter = 0
        tbl_counter = 0
        
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                
                if para:
                    para_text = ""
                    # Check for inline images in runs
                    for run in para.runs:
                        if run._element.xpath('.//a:blip'):
                            # Found an image
                            img_counter += 1
                            anchor = f"##IMG{img_counter:03d}##"
                            para_text += anchor
                            
                            # Extract image data
                            try:
                                for rel in run._element.xpath('.//a:blip/@r:embed'):
                                    image_part = doc.part.related_parts[rel]
                                    objects[anchor] = {
                                        'type': 'image',
                                        'data': image_part.blob,
                                        'content_type': image_part.content_type
                                    }
                            except Exception as e:
                                print(f"Error extracting image: {e}")
                        else:
                            para_text += run.text
                    
                    if para_text.strip():
                        text_parts.append(para_text)
            
            # Handle tables
            elif element.tag.endswith('tbl'):
                tbl_counter += 1
                anchor = f"##TBL{tbl_counter:03d}##"
                text_parts.append(anchor)
                
                # Find the table object
                for table in doc.tables:
                    if table._element == element:
                        # Store table structure
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text for cell in row.cells]
                            table_data.append(row_data)
                        
                        objects[anchor] = {
                            'type': 'table',
                            'data': table_data,
                            'table_element': table._element
                        }
                        break
        
        print(f"[DEBUG] DOCX Parser: Found {img_counter} images, {tbl_counter} tables")
        print(f"[DEBUG] Objects dict keys: {list(objects.keys())}")
        full_text = "\n".join(text_parts)
        print(f"[DEBUG] Text with anchors (first 500 chars): {full_text[:500]}")
        
        return full_text, objects

    def _parse_html(self, content: bytes) -> str:
        soup = BeautifulSoup(content, "html.parser")
        return soup.get_text(separator="\n")

class FileExporter:
    async def export_file(self, text: str, filename: str, content_type: str, objects: dict = None) -> bytes:
        """Export translated text to the same format as the original file"""
        filename_lower = filename.lower()
        objects = objects or {}
        
        if filename_lower.endswith(".docx"):
            return self._export_docx(text, objects)
        elif filename_lower.endswith(".txt"):
            return text.encode("utf-8")
        elif filename_lower.endswith(".html"):
            return self._export_html(text)
        elif filename_lower.endswith(".pdf"):
            # Export PDF as DOCX to preserve objects
            return self._export_docx(text, objects)
        else:
            return text.encode("utf-8")
    
    def _export_docx(self, text: str, objects: dict = None) -> bytes:
        """Export text to DOCX format with objects"""
        doc = docx.Document()
        objects = objects or {}
        
        import re
        # Pattern to match anchors: ##IMG001##, ##TBL001##, etc.
        anchor_pattern = r'##(IMG|TBL)(\d{3})##'
        
        for line in text.split("\n"):
            # Check if line contains anchors
            matches = list(re.finditer(anchor_pattern, line))
            
            if matches:
                # Process line with anchors
                last_end = 0
                para = doc.add_paragraph()
                
                for match in matches:
                    # Add text before anchor
                    if match.start() > last_end:
                        para.add_run(line[last_end:match.start()])
                    
                    # Handle anchor
                    anchor = match.group(0)
                    if anchor in objects:
                        obj = objects[anchor]
                        
                        if obj['type'] == 'image':
                            # Add image
                            try:
                                img_stream = io.BytesIO(obj['data'])
                                run = para.add_run()
                                run.add_picture(img_stream)
                            except Exception as e:
                                print(f"Error adding image: {e}")
                                para.add_run(f"[Image: {anchor}]")
                        
                        elif obj['type'] == 'table':
                            # Tables need to be added separately, not inline
                            # Add table after current paragraph
                            table_data = obj['data']
                            if table_data:
                                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                                for i, row_data in enumerate(table_data):
                                    for j, cell_text in enumerate(row_data):
                                        table.rows[i].cells[j].text = cell_text
                    
                    last_end = match.end()
                
                # Add remaining text after last anchor
                if last_end < len(line):
                    para.add_run(line[last_end:])
            
            elif line.strip():
                # Regular paragraph without anchors
                doc.add_paragraph(line)
        
        print(f"[DEBUG] DOCX Exporter: Received {len(objects)} objects")
        print(f"[DEBUG] Objects keys: {list(objects.keys())}")
        print(f"[DEBUG] Text to export (first 500 chars): {text[:500]}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _export_html(self, text: str) -> bytes:
        """Export text to HTML format"""
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Translated Document</title>
</head>
<body>
"""
        for paragraph in text.split("\n"):
            if paragraph.strip():
                html += f"    <p>{paragraph}</p>\n"
        html += """</body>
</html>"""
        return html.encode("utf-8")
    
    def _export_pdf(self, text: str) -> bytes:
        """Export text to PDF format using reportlab"""
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        style_normal = styles['Normal']
        style_normal.fontSize = 11
        style_normal.leading = 14
        style_normal.alignment = TA_LEFT
        
        # Add paragraphs
        for line in text.split('\n'):
            if line.strip():
                # Create paragraph and add to elements
                para = Paragraph(line.strip(), style_normal)
                elements.append(para)
                elements.append(Spacer(1, 0.2 * inch))
            else:
                # Add extra space for empty lines
                elements.append(Spacer(1, 0.3 * inch))
        
        # Build PDF
        doc.build(elements)
        
        # Get PDF bytes
        buffer.seek(0)
        return buffer.read()
